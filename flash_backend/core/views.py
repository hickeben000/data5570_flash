import io

from django.conf import settings
from django.contrib.auth import authenticate
from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import status, viewsets
from rest_framework.authtoken.models import Token
from rest_framework.parsers import FormParser, JSONParser, MultiPartParser
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from .ai_utils import (
    AIConfigurationError,
    AIProviderError,
    AIResponseError,
    generate_flashcards,
    generate_quiz,
    grade_free_response_batch,
)
from .models import (
    AnswerChoice,
    Course,
    Document,
    Flashcard,
    FlashcardDeck,
    Quiz,
    QuizQuestion,
)
from .serializers import (
    CourseSerializer,
    DocumentSerializer,
    FlashcardDeckSerializer,
    FlashcardSerializer,
    QuizResultSerializer,
    QuizSerializer,
    UserSerializer,
)


# ──────────────────────────── Helpers ─────────────────────────


def _check_magic_bytes(file, name: str) -> None:
    """
    Validate file format against magic bytes.
    Reads the first 4 bytes, then seeks back to 0 so callers see the full file.
    Raises ValueError with a user-facing message on mismatch.
    """
    header = file.read(4)
    file.seek(0)
    if name.endswith(".pdf") and not header.startswith(b"%PDF"):
        raise ValueError(
            "The uploaded file does not appear to be a valid PDF."
        )
    if name.endswith(".docx") and not header.startswith(b"PK"):
        raise ValueError(
            "The uploaded file does not appear to be a valid DOCX file."
        )


def _truncate_text(text: str) -> str:
    limit = getattr(settings, "MAX_DOCUMENT_CHARS", 100_000)
    return text[:limit] if len(text) > limit else text


def extract_text_from_file(file) -> str:
    """
    Extract plain text from an uploaded file.
    Supports: PDF, DOCX, and plain text (TXT / fallback).

    TODO (enhancement): add support for .pptx (python-pptx) and images (pytesseract OCR)
                        if students upload slides or scanned notes.
    TODO (enhancement): move extraction + AI calls to a background task queue (e.g. Celery)
                        to avoid blocking the request thread on large files.
    """
    name = file.name.lower()
    _check_magic_bytes(file, name)

    if name.endswith(".pdf"):
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file.read()))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if not text.strip():
            raise ValueError(
                "No text could be extracted from this PDF. "
                "It may be a scanned image — try a text-based PDF."
            )
        return text

    elif name.endswith(".docx"):
        import docx
        doc = docx.Document(io.BytesIO(file.read()))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    else:
        # Fall back to reading as UTF-8 text (handles .txt and similar).
        return file.read().decode("utf-8", errors="replace")


def _validate_upload_size(file) -> None:
    """Raise ValueError if the uploaded file exceeds MAX_UPLOAD_BYTES."""
    limit = getattr(settings, "MAX_UPLOAD_BYTES", 20 * 1024 * 1024)
    if file.size > limit:
        mb = limit // (1024 * 1024)
        raise ValueError(f"File is too large. Maximum allowed size is {mb} MB.")


def get_course_for_user(course_id, user):
    """404 if missing or not owned by user."""
    return get_object_or_404(Course, pk=course_id, user=user)


def get_document_for_user(document_id, user):
    """404 if missing or not owned by user (via course)."""
    return get_object_or_404(Document, pk=document_id, course__user=user)


def get_deck_for_user(pk, user):
    """404 if missing or not owned by user."""
    return get_object_or_404(FlashcardDeck, pk=pk, document__course__user=user)


def get_quiz_for_user(pk, user):
    """404 if missing or not owned by user."""
    return get_object_or_404(Quiz, pk=pk, document__course__user=user)


def get_flashcard_for_user(pk, user):
    """404 if missing or not owned by user."""
    return get_object_or_404(Flashcard, pk=pk, deck__document__course__user=user)


def parse_int_field(value, *, default, field_name, minimum=0):
    if value in (None, ""):
        return default
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        raise ValueError(f"{field_name} must be an integer.")
    if parsed < minimum:
        raise ValueError(f"{field_name} must be at least {minimum}.")
    return parsed


def get_request_api_key(request):
    return (request.headers.get("X-OpenAI-Api-Key") or "").strip()


def request_uses_https(request):
    if request.is_secure():
        return True
    return request.headers.get("X-Forwarded-Proto", "").lower() == "https"


def validate_ai_request_security(request):
    if (
        get_request_api_key(request)
        and getattr(settings, "REQUIRE_HTTPS_FOR_AI", False)
        and not request_uses_https(request)
    ):
        raise AIConfigurationError(
            "AI requests must use HTTPS in deployed environments."
        )


def ai_error_response(exc):
    if isinstance(exc, AIProviderError):
        return Response({"error": str(exc)}, status=status.HTTP_502_BAD_GATEWAY)
    return Response({"error": str(exc)}, status=status.HTTP_400_BAD_REQUEST)


# ──────────────────────────── Auth ────────────────────────────


class RegisterView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        serializer = UserSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class LoginView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        username = request.data.get("username")
        password = request.data.get("password")
        user = authenticate(username=username, password=password)
        if user is not None:
            # get_or_create ensures repeated logins return the same token.
            token, _ = Token.objects.get_or_create(user=user)
            return Response({
                "token": token.key,
                "user_id": user.id,
                "username": user.username,
            })
        return Response(
            {"error": "Invalid credentials"},
            status=status.HTTP_401_UNAUTHORIZED,
        )


class HealthCheckView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        return Response(
            {
                "status": "ok",
                "debug": settings.DEBUG,
            }
        )


# ──────────────────────────── Courses ─────────────────────────


class CourseViewSet(viewsets.ModelViewSet):
    serializer_class = CourseSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        # Only return courses belonging to the authenticated user.
        return Course.objects.filter(user=self.request.user)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)


# ──────────────────────────── Documents ───────────────────────


class DocumentListCreateView(APIView):
    """
    GET  /api/courses/<course_id>/documents/  — list documents in a course
    POST /api/courses/<course_id>/documents/  — create a document

    Accepts either:
      - JSON body (application/json) with a `raw_text` field (source_type="paste")
      - Multipart form (multipart/form-data) with a `file` field (source_type="upload");
        text is extracted server-side
    """
    parser_classes = [JSONParser, MultiPartParser, FormParser]

    def get(self, request, course_id):
        get_course_for_user(course_id, request.user)
        documents = Document.objects.filter(
            course_id=course_id, course__user=request.user
        )
        serializer = DocumentSerializer(documents, many=True)
        return Response(serializer.data)

    def post(self, request, course_id=None):
        get_course_for_user(course_id, request.user)
        data = request.data.copy()
        if course_id:
            data["course"] = course_id

        uploaded_file = request.FILES.get("file")
        if uploaded_file:
            try:
                _validate_upload_size(uploaded_file)
                raw_text = extract_text_from_file(uploaded_file)
            except ValueError as exc:
                return Response({"error": str(exc)}, status=status.HTTP_400_BAD_REQUEST)
            data["raw_text"] = _truncate_text(raw_text)
            data["source_type"] = "upload"
            # Use the original filename as the title if none was provided.
            if not data.get("title"):
                data["title"] = uploaded_file.name
        elif data.get("raw_text", "").strip():
            data["raw_text"] = _truncate_text(data["raw_text"])
            if not data.get("source_type"):
                data["source_type"] = "paste"
        else:
            return Response(
                {"error": "Provide either a file upload or pasted text."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        serializer = DocumentSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class DocumentCreateView(APIView):
    """
    POST /api/documents/ — standalone document creation (course ID in request body).
    Same file-upload handling as DocumentListCreateView.
    """
    parser_classes = [JSONParser, MultiPartParser, FormParser]

    def post(self, request):
        data = request.data.copy()
        course_pk = data.get("course")
        if course_pk is None:
            return Response(
                {"course": ["This field is required."]},
                status=status.HTTP_400_BAD_REQUEST,
            )
        get_course_for_user(course_pk, request.user)

        uploaded_file = request.FILES.get("file")
        if uploaded_file:
            try:
                _validate_upload_size(uploaded_file)
                raw_text = extract_text_from_file(uploaded_file)
            except ValueError as exc:
                return Response({"error": str(exc)}, status=status.HTTP_400_BAD_REQUEST)
            data["raw_text"] = _truncate_text(raw_text)
            data["source_type"] = "upload"
            if not data.get("title"):
                data["title"] = uploaded_file.name
        elif data.get("raw_text", "").strip():
            data["raw_text"] = _truncate_text(data["raw_text"])
            if not data.get("source_type"):
                data["source_type"] = "paste"
        else:
            return Response(
                {"error": "Provide either a file upload or pasted text."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        serializer = DocumentSerializer(data=data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# ──────────────────────────── Flashcards ──────────────────────


def _combine_document_texts(primary_doc, additional_ids, user) -> str:
    """
    Build combined raw_text from primary_doc plus any additional document IDs.
    Each additional document is ownership-checked before its text is included.
    The result is truncated to MAX_DOCUMENT_CHARS.
    """
    parts = [primary_doc.raw_text]
    for doc_id in additional_ids:
        extra_doc = get_document_for_user(doc_id, user)
        parts.append(extra_doc.raw_text)
    return _truncate_text("\n\n".join(parts))


class GenerateFlashcardsView(APIView):
    def post(self, request, document_id):
        document = get_document_for_user(document_id, request.user)
        additional_ids = request.data.get("additional_document_ids") or []
        try:
            validate_ai_request_security(request)
            num_cards = parse_int_field(
                request.data.get("num_cards"),
                default=10,
                field_name="num_cards",
                minimum=1,
            )
            extra_prompt = request.data.get("extra_prompt", "")
            combined_text = _combine_document_texts(document, additional_ids, request.user)
            card_data = generate_flashcards(
                combined_text,
                num_cards=num_cards,
                extra_prompt=extra_prompt,
                api_key=get_request_api_key(request),
            )
        except (AIConfigurationError, AIProviderError, AIResponseError) as exc:
            return ai_error_response(exc)
        except ValueError as exc:
            return Response({"error": str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        deck = FlashcardDeck.objects.create(
            document=document,
            title=f"Flashcards — {document.title}",
        )
        for card in card_data:
            Flashcard.objects.create(
                deck=deck, front=card["front"], back=card["back"]
            )

        serializer = FlashcardDeckSerializer(deck)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class FlashcardDeckDetailView(APIView):
    def get(self, request, pk):
        deck = get_deck_for_user(pk, request.user)
        serializer = FlashcardDeckSerializer(deck)
        return Response(serializer.data)


class FlashcardUpdateView(APIView):
    def put(self, request, pk):
        card = get_flashcard_for_user(pk, request.user)
        serializer = FlashcardSerializer(card, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


# ──────────────────────────── Quizzes ─────────────────────────


class GenerateQuizView(APIView):
    def post(self, request, document_id):
        document = get_document_for_user(document_id, request.user)
        additional_ids = request.data.get("additional_document_ids") or []
        try:
            validate_ai_request_security(request)
            difficulty = request.data.get("difficulty", "medium")
            mc_count = parse_int_field(
                request.data.get("mc_count"),
                default=2,
                field_name="mc_count",
                minimum=0,
            )
            fitb_count = parse_int_field(
                request.data.get("fitb_count"),
                default=1,
                field_name="fitb_count",
                minimum=0,
            )
            fr_count = parse_int_field(
                request.data.get("fr_count"),
                default=1,
                field_name="fr_count",
                minimum=0,
            )
            if mc_count + fitb_count + fr_count <= 0:
                raise ValueError("At least one quiz question is required.")

            class_name = request.data.get("class_name", "")
            learning_objectives = request.data.get("learning_objectives", "")
            extra_prompt = request.data.get("extra_prompt", "")
            combined_text = _combine_document_texts(document, additional_ids, request.user)

            question_data = generate_quiz(
                combined_text,
                difficulty,
                mc_count,
                fitb_count,
                fr_count,
                class_name=class_name,
                learning_objectives=learning_objectives,
                extra_prompt=extra_prompt,
                api_key=get_request_api_key(request),
            )
        except (AIConfigurationError, AIProviderError, AIResponseError) as exc:
            return ai_error_response(exc)
        except ValueError as exc:
            return Response({"error": str(exc)}, status=status.HTTP_400_BAD_REQUEST)

        quiz = Quiz.objects.create(
            document=document,
            difficulty=difficulty,
            class_name=class_name,
            learning_objectives=learning_objectives,
        )

        for q in question_data:
            question = QuizQuestion.objects.create(
                quiz=quiz,
                question_type=q["question_type"],
                question_text=q["question_text"],
            )
            # Create an AnswerChoice row for each option.
            # MC: multiple choices, one is_correct=True.
            # FITB: single choice row with the correct answer, is_correct=True.
            # Free response: no choices created here — graded by Gemini after submission.
            for choice in q.get("answer_choices", []):
                AnswerChoice.objects.create(
                    question=question,
                    choice_text=choice["choice_text"],
                    is_correct=choice["is_correct"],
                )

        serializer = QuizSerializer(quiz)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class QuizDetailView(APIView):
    def get(self, request, pk):
        quiz = get_quiz_for_user(pk, request.user)
        serializer_class = QuizResultSerializer if quiz.completed_at else QuizSerializer
        serializer = serializer_class(quiz)
        return Response(serializer.data)


class QuizSubmitView(APIView):
    def put(self, request, pk):
        quiz = get_quiz_for_user(pk, request.user)

        # Expected format: {"answers": {"<question_id>": "<value>"}}
        # MC:            value = AnswerChoice ID (as string), e.g. "7"
        # FITB:          value = user's typed text,           e.g. "mitosis"
        # free_response: value = user's full written answer
        answers = request.data.get("answers", {})
        questions = quiz.questions.all()
        total = questions.count()
        auto_correct = 0

        free_response_questions = []
        try:
            if questions.filter(question_type="free_response").exists():
                validate_ai_request_security(request)
        except AIConfigurationError as exc:
            return ai_error_response(exc)

        for question in questions:
            user_answer = answers.get(str(question.id), "")
            question.user_answer = user_answer

            if question.question_type == "mc":
                # Look up the chosen AnswerChoice and check its is_correct flag.
                try:
                    chosen = question.answer_choices.get(id=int(user_answer))
                    question.is_correct = chosen.is_correct
                except (AnswerChoice.DoesNotExist, ValueError, TypeError):
                    question.is_correct = False
                if question.is_correct:
                    auto_correct += 1
                question.save()

            elif question.question_type == "fitb":
                # Compare typed answer to the correct AnswerChoice text (case-insensitive).
                correct_choice = question.answer_choices.filter(is_correct=True).first()
                if correct_choice:
                    question.is_correct = (
                        user_answer.strip().lower()
                        == correct_choice.choice_text.strip().lower()
                    )
                    if question.is_correct:
                        question.explanation = (
                            f'Correct! "{correct_choice.choice_text}" is the right answer.'
                        )
                    else:
                        question.explanation = (
                            f'The correct answer is "{correct_choice.choice_text}". '
                            f'You answered "{user_answer.strip() or "(blank)"}".'
                        )
                else:
                    question.is_correct = False
                    question.explanation = "No correct answer was stored for this question."
                if question.is_correct:
                    auto_correct += 1
                question.save()

            elif question.question_type == "free_response":
                # Free response is graded by Gemini after collecting all answers.
                # Save the answer now; is_correct/feedback set below.
                question.save()
                free_response_questions.append(question)

        if free_response_questions:
            try:
                results = grade_free_response_batch(
                    free_response_questions,
                    document_text=quiz.document.raw_text,
                    class_name=quiz.class_name,
                    learning_objectives=quiz.learning_objectives,
                    api_key=get_request_api_key(request),
                )
            except (AIConfigurationError, AIProviderError, AIResponseError) as exc:
                return ai_error_response(exc)
            for question, result in zip(free_response_questions, results):
                question.is_correct = result.get("is_correct", False)
                question.feedback = result.get("feedback", "")
                question.explanation = result.get("explanation", "")
                question.save()
                if question.is_correct:
                    auto_correct += 1

        quiz.score = (auto_correct / total * 100) if total > 0 else 0
        quiz.completed_at = timezone.now()
        quiz.save()

        # Return the full result serializer so the frontend gets graded answers.
        serializer = QuizResultSerializer(quiz)
        return Response(serializer.data)
